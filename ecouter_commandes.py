import firebase_admin
from firebase_admin import credentials, db
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from docx import Document
from openpyxl import Workbook
from datetime import datetime
import time
import os

# ========== CONFIG ==========
FORMAT_FACTURE = "pdf"  # Mets "pdf" ou "word" ou "excel" - un seul à la fois
URL_BASE_DONNEES = "https://sunu-com-default-rtdb.firebaseio.com/"
# ============================
# 1. CONNECTE FIREBASE - VERSION RENDER
firebase_json = os.environ.get('FIREBASE_KEY_JSON')
if not firebase_json:
    print("❌ STOP : Variable FIREBASE_KEY_JSON manquante sur Render")
    exit()

cred = credentials.Certificate(json.loads(firebase_json))
firebase_admin.initialize_app(cred, {'databaseURL': URL_BASE_DONNEES})

# 1. CONNECTE FIREBASE
cred = credentials.Certificate(CHEMIN_CLE_JSON)
firebase_admin.initialize_app(cred, {
    'databaseURL': URL_BASE_DONNEES
})

def generer_pdf(data, order_id):
    nom_fichier = f"Facture_SUNU_{order_id}_{int(time.time())}.pdf"
    c = canvas.Canvas(nom_fichier, pagesize=A4)
    width, height = A4
    
    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, height-50, "SUNU-COM - FACTURE")
    c.setFont("Helvetica", 10)
    c.drawString(50, height-70, f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    c.setFont("Helvetica", 12)
    y = height - 120
    c.drawString(50, y, f"Commande N°: {order_id[:8]}")
    y -= 25
    c.drawString(50, y, f"Client: {data.get('nom', 'Client')}")
    y -= 20
    c.drawString(50, y, f"Tel: {data.get('telephone', 'N/A')}")
    y -= 20
    c.drawString(50, y, f"Adresse: {data.get('adresse', 'Non renseignée')}")
    
    y -= 40
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "DÉTAILS COMMANDE:")
    y -= 25
    
    c.setFont("Helvetica", 10)
    panier = data.get('panier', {})
    total_calcule = 0
    
    if isinstance(panier, dict) and panier:
        for key, item in panier.items():
            if isinstance(item, dict):
                nom_item = item.get('nom', 'Produit')
                qty = int(item.get('qty', 1))
                prix = int(item.get('prix', 0))
                sous_total = qty * prix
                total_calcule += sous_total
                c.drawString(70, y, f"- {nom_item} x{qty} = {sous_total:,} FCFA".replace(",", " "))
                y -= 20
    else:
        total = int(data.get('total', 0))
        c.drawString(70, y, f"- Commande diverse = {total:,} FCFA".replace(",", " "))
        total_calcule = total
        y -= 20
    
    y -= 15
    c.line(50, y, 250, y)
    y -= 25
    c.setFont("Helvetica-Bold", 14)
    total_final = data.get('total', total_calcule)
    c.drawString(50, y, f"TOTAL: {total_final:,} FCFA".replace(",", " "))
    
    c.save()
    return nom_fichier

def generer_word(data, order_id):
    nom_fichier = f"Facture_SUNU_{order_id}_{int(time.time())}.docx"
    doc = Document()
    
    doc.add_heading('SUNU-COM - FACTURE', 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph(f"Commande N°: {order_id[:8]}")
    doc.add_paragraph(f"Client: {data.get('nom', 'Client')}")
    doc.add_paragraph(f"Tel: {data.get('telephone', 'N/A')}")
    doc.add_paragraph(f"Adresse: {data.get('adresse', 'Non renseignée')}")
    
    doc.add_heading('DÉTAILS COMMANDE:', level=1)
    panier = data.get('panier', {})
    total_calcule = 0
    
    if isinstance(panier, dict) and panier:
        for key, item in panier.items():
            if isinstance(item, dict):
                nom_item = item.get('nom', 'Produit')
                qty = int(item.get('qty', 1))
                prix = int(item.get('prix', 0))
                sous_total = qty * prix
                total_calcule += sous_total
                doc.add_paragraph(f"- {nom_item} x{qty} = {sous_total:,} FCFA".replace(",", " "))
    else:
        total = int(data.get('total', 0))
        doc.add_paragraph(f"- Commande diverse = {total:,} FCFA".replace(",", " "))
        total_calcule = total
    
    doc.add_paragraph("---")
    total_final = data.get('total', total_calcule)
    doc.add_paragraph(f"TOTAL: {total_final:,} FCFA".replace(",", " "))
    
    doc.save(nom_fichier)
    return nom_fichier

def generer_excel(data, order_id):
    nom_fichier = f"Facture_SUNU_{order_id}_{int(time.time())}.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Facture"
    
    ws['A1'] = 'SUNU-COM - FACTURE'
    ws['A2'] = f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws['A4'] = f"Commande N°: {order_id[:8]}"
    ws['A5'] = f"Client: {data.get('nom', 'Client')}"
    ws['A6'] = f"Tel: {data.get('telephone', 'N/A')}"
    ws['A7'] = f"Adresse: {data.get('adresse', 'Non renseignée')}"
    
    ws['A9'] = 'PRODUIT'
    ws['B9'] = 'QTÉ'
    ws['C9'] = 'PRIX'
    ws['D9'] = 'TOTAL'
    
    row = 10
    panier = data.get('panier', {})
    total_calcule = 0
    
    if isinstance(panier, dict) and panier:
        for key, item in panier.items():
            if isinstance(item, dict):
                nom_item = item.get('nom', 'Produit')
                qty = int(item.get('qty', 1))
                prix = int(item.get('prix', 0))
                sous_total = qty * prix
                total_calcule += sous_total
                ws[f'A{row}'] = nom_item
                ws[f'B{row}'] = qty
                ws[f'C{row}'] = prix
                ws[f'D{row}'] = sous_total
                row += 1
    else:
        total = int(data.get('total', 0))
        ws[f'A{row}'] = 'Commande diverse'
        ws[f'D{row}'] = total
        total_calcule = total
    
    ws[f'A{row+1}'] = 'TOTAL:'
    ws[f'D{row+1}'] = data.get('total', total_calcule)
    
    wb.save(nom_fichier)
    return nom_fichier

def generer_facture(data, order_id):
    try:
        if FORMAT_FACTURE == "word":
            nom_fichier = generer_word(data, order_id)
        elif FORMAT_FACTURE == "excel":
            nom_fichier = generer_excel(data, order_id)
        else:
            nom_fichier = generer_pdf(data, order_id)
        
        print(f"✅ Facture générée: {nom_fichier}")
        return nom_fichier
        
    except Exception as e:
        print(f"❌ ERREUR génération: {e}")
        return None

def ecouter():
    ref = db.reference('commandes')
    
    def callback(event):
        try:
            path = event.path
            data = event.data
            
            if data is None:
                return
            
            if isinstance(data, dict) and path == '/':
                for order_id, commande_data in data.items():
                    if isinstance(commande_data, dict) and commande_data.get('status') == 'pending':
                        print(f"\n🔔 Nouvelle commande de {commande_data.get('nom', 'Client')} - {commande_data.get('total', 0)} FCFA")
                        pdf = generer_facture(commande_data, order_id)
                        if pdf:
                            ref.child(order_id).update({'status': 'traitee'})
                            print(f"📄 Status mis à jour: traitee")
            
            elif isinstance(data, dict) and '/' not in path[1:] and data.get('status') == 'pending':
                order_id = path.split('/')[-1]
                print(f"\n🔔 Nouvelle commande de {data.get('nom', 'Client')} - {data.get('total', 0)} FCFA")
                pdf = generer_facture(data, order_id)
                if pdf:
                    ref.child(order_id).update({'status': 'traitee'})
                    print(f"📄 Status mis à jour: traitee")
                
        except Exception as e:
            print(f"❌ Erreur callback: {e}")
    
    ref.listen(callback)
    print("👂 En écoute des nouvelles commandes... Ne ferme pas cette fenêtre")
    while True:
        time.sleep(1)

if __name__ == '__main__':
    print(f"🧪 TEST DIRECT - Format: {FORMAT_FACTURE.upper()}")
    
    commande_test = {
        "nom": "Client Test",
        "telephone": "701234567",
        "adresse": "Plateau Dakar",
        "total": 7500,
        "panier": {
            "0": {"nom": "Burger Poulet", "qty": 1, "prix": 5000},
            "1": {"nom": "Frites", "qty": 1, "prix": 2500}
        },
        "status": "pending"
    }
    
    pdf = generer_facture(commande_test, "TEST123")
    if pdf:
        print("🎉 Test OK ! La facture est créée")
    else:
        print("💥 Test KO")
